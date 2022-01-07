import docx
doc = docx.Document('testData.txt')
docdoc.paragraphs
len(doc.paragraphs)
doc.paragraphs[0].text
for para in doc.paragraphs:
    print(para.text)

doc.paragraphs
doc.para[0].runs
len.doc.para[0].runs
len.doc.para[0].runs[0]
len.doc.para[0].runs[0].text
len.doc.para[0].runs[1].text
len.doc.para[0].runs[2].text
len.doc.para[0].runs[3].text
len.doc.para[0].runs[4].text

doc = docxDocument()
doc.add_paragraph('Hello World')
paraObject = doc.add_paragraph('This is a lot more text to our document')
paraObject2 = doc.add_paragraph('This is even more content')
paraObject.add_run('Im adding more text')
doc.save('newDoc.docx')

doc = docx.Document('newDoc.docx')
doc.paragraphs[0].style
doc.paragraphs[0].runs[0].bold = True
doc.paragraphs[0].runs[0].underline = True
doc.save('newDoc.docx')

doc = docx.Document('newDoc.docx')
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)

doc = docx.Document('newDoc.docx')
doc.add_picture('Belts.png', width=docx.shared.Inches(2), height=docxshared.Inches(2))
docsave('newDox.docx')